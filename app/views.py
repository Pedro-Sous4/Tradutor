from django.http import HttpResponse
from django.shortcuts import render
from project import settings
from django.core.files.storage import FileSystemStorage
import pandas as pd
from docx import Document
from django.contrib import messages
import os
from django.conf import settings



def upload_file(request):
    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']

        # Verificar se o arquivo enviado é .docx
        if not uploaded_file.name.endswith('.docx'):
            messages.error(request, 'Por favor, envie um arquivo .docx.')
            return render(request, 'upload_form.html')

        # Armazenar o arquivo temporariamente
        fs = FileSystemStorage()
        filename = fs.save(uploaded_file.name, uploaded_file)
        file_path = fs.path(filename)

        try:
            # Ler o dicionário de variáveis do Excel
            dict_file = os.path.join(settings.BASE_DIR, 'dicionario.xlsx')  # Caminho dinâmico do dicionário
            dicionario = pd.read_excel(dict_file)

            # Garantir que as colunas esperadas estão no arquivo
            if 'esolution' not in dicionario.columns or 'sienge' not in dicionario.columns:
                messages.error(request, 'O arquivo dicionário está incorreto. Verifique as colunas.')
                return render(request, 'upload_form.html')

            # Abrir o documento .docx
            doc = Document(file_path)

            # Função auxiliar para substituir texto dentro dos runs
            def substituir_texto(texto):
                for esolution, sienge in zip(dicionario['esolution'], dicionario['sienge']):
                    texto = texto.replace(esolution, sienge)
                return texto

            # Função para processar runs e manter formatação
            def processar_paragrafo(paragraph):
                for run in paragraph.runs:
                    run.text = substituir_texto(run.text)

            # Substituir variáveis no documento - parágrafos e títulos
            for paragraph in doc.paragraphs:
                processar_paragrafo(paragraph)

            # Substituir variáveis dentro das tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            processar_paragrafo(paragraph)

            # Salvar o documento processado
            new_filename = 'processed_' + filename
            new_file_path = fs.path(new_filename)
            doc.save(new_file_path)

            # Forçar o download do arquivo processado
            with open(new_file_path, 'rb') as file:
                response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename={new_filename}'
                return response

        except Exception as e:
            # Tratamento de exceções
            messages.error(request, f'Ocorreu um erro ao processar o arquivo: {str(e)}')
            return render(request, 'upload_form.html')

    # Se o metodo não for POST ou não houver arquivo, exibe o formulário de upload
    return render(request, 'upload_form.html')